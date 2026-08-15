import os
from pathlib import Path
from datetime import datetime

from docx import Document
import pymupdf # PyMuPDF

import re

# def read_file(filepath: str) -> dict:
#     """
#     Reads a TXT, DOCX, or PDF file and returns its content along with metadata.

#     Args:
#         filepath (str): Path to the file.

#     Returns:
#         dict: Dictionary containing success status, metadata, and extracted content.
#     """
#     try:
#         path  = Path(filepath)

#         if not path.exists():
#             return {
#                 "success": False,
#                 "message": f"File '{filepath}' does not exist.",
#                 "data": None
#             }

#         metadata = {
#             "filepath": str(path),
#             "filename": path.name,
#             "extension": path.suffix.lower(),
#             "size": path.stat().st_size, 
#             "modified_date": datetime.fromtimestamp(path.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")
#         }

#         if path.suffix.lower() == ".txt":
#             encodings = ["utf-8", "utf-8-sig", "cp1252", "latin-1"]

#             content = None

#             for encoding in encodings:
#                 try:
#                     with open(path, "r", encoding=encoding) as f:
#                         content = f.read()
#                     break
#                 except UnicodeDecodeError:
#                     continue

#             if content is None:
#                 raise UnicodeDecodeError(...)
            
#         elif path.suffix.lower() == ".docx":
#             document = Document(path)
#             content = "\n".join(paragraph.text for paragraph in document.paragraphs)
#         elif path.suffix.lower() == ".pdf":
#             document = pymupdf.open(path)

#             content = ""

#             for page in document:
#                 content += page.get_text()

#             document.close()
#         else:
#             return {
#                 "success": False,
#                 "message": f"Unsupported file type: {path.suffix}",
#                 "data": None
#             }

#         # return {
#         #     "success": True, 
#         #     **metadata,
#         #     "content": content
#         # }
#         return {
#             "success": True,
#             "message": "File read successfully.",
#             "data": {
#                 **metadata,
#                 "content": content
#             }
#         }
#     except Exception as e:
#         return {
#             "success": False,
#             "message": str(e),
#             "data": None
#         }

def read_file(filepath: str) -> dict:
    """
    Reads a TXT, DOCX, or PDF file and returns its content along with metadata.

    Args:
        filepath (str): Path to the file.

    Returns:
        dict: Dictionary containing success status, metadata, and extracted content.
    """

    try:
        path = Path(filepath)

        if not path.exists():
            return {
                "success": False,
                "message": f"File '{filepath}' does not exist.",
                "data": None
            }

        extension = path.suffix.lower()

        metadata = {
            "filepath": str(path),
            "filename": path.name,
            "extension": extension,
            "size": path.stat().st_size,
            "modified_date": datetime.fromtimestamp(
                path.stat().st_mtime
            ).strftime("%Y-%m-%d %H:%M:%S")
        }

        # ---------------- TXT ---------------- #
        if extension == ".txt":

            encodings = [
                "utf-8",
                "utf-8-sig",
                "cp1252",
                "latin-1"
            ]

            content = None

            for encoding in encodings:
                try:
                    with path.open("r", encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise Exception(
                    "Unable to decode the text file. Unsupported or corrupted encoding."
                )

        # ---------------- DOCX ---------------- #
        elif extension == ".docx":

            document = Document(path)
            content = "\n".join(
                paragraph.text for paragraph in document.paragraphs
            )

        # ---------------- PDF ---------------- #
        elif extension == ".pdf":

            content = ""

            with pymupdf.open(path) as document:
                for page in document:
                    content += page.get_text()

        # ---------------- Unsupported ---------------- #
        else:
            return {
                "success": False,
                "message": f"Unsupported file type: {extension}",
                "data": None
            }

        return {
            "success": True,
            "message": "File read successfully.",
            "data": {
                **metadata,
                "content": content
            }
        }

    except Exception as e:
        return {
            "success": False,
            "message": str(e),
            "data": None
        }


def list_files(directory: str, extension: str = None) -> list:
    """
    Lists files in a directory with optional extension filtering.

    Args:
        directory (str): Directory path.
        extension (str, optional): Filter by extension (e.g. ".pdf" or "pdf").

    Returns:
        list: List of dictionaries containing metadata for each file.
    """
    path = Path(directory)

    if not path.exists():
        return {
            "success": False,
            "message": f"Directory '{directory}' does not exist.",
            "data": None
        }

    if not path.is_dir():
        return {
            "success": False,
            "message": f"'{directory}' is not a directory.",
            "data": None
        }
         
    
    if extension:
        extension = extension.lower()

        if not extension.startswith("."):
            extension = "." + extension

    files = []

    for file in path.iterdir():
        if not file.is_file():
            continue
        if extension and file.suffix.lower() != extension:
            continue

        files.append(
            {
                "filename": file.name,
                "filepath": str(file),
                "extension": file.suffix.lower(),
                "size": file.stat().st_size,
                "modified_date": datetime.fromtimestamp(file.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")
            }
        )

    return {
        "success": True,
        "message": "Files listed successfully.",
        "data": {
            "count": len(files),
            "files": files
        }
    }


def write_file(filepath: str, content: str) -> dict:
    """
    Writes text content to a file.
    Creates parent directories if they do not exist.

    Args:
        filepath (str): Destination file path.
        content (str): Content to write.

    Returns:
        dict: Success status, message, and file metadata.
    """
    try:
        path = Path(filepath)

        #Create Parent Directories(the parameters 'parents'==> create every parent dir if not present
        # & 'exist_ok' ==> if exist continue)
        path.parent.mkdir(parents=True, exist_ok=True)

        with open(path, "w", encoding="utf-8") as file:
            file.write(content)

        metadata = {
            "filepath": str(path),
            "filename": path.name,
            "extension": path.suffix.lower(),
            "size": path.stat().st_size,
            "modified_date": datetime.fromtimestamp(
                path.stat().st_mtime
            ).strftime("%Y-%m-%d %H:%M:%S")
        }

        return {
            "success": True,
            "message": "File written successfully.",
            "data": metadata
        }
    except Exception as e:
        return {
            "success": False,
            "message": str(e),
            "data": None
        }


def search_in_file(filepath: str, keyword: str) -> dict:
    """
    Searches for a keyword in a file (TXT, DOCX, PDF).

    Args:
        filepath (str): Path to the file.
        keyword (str): Keyword to search.

    Returns:
        dict: Search results with matching lines.
    """
    try:
        if not keyword.strip():
            return {
                "success": False,
                "message": "Keyword cannot be empty.",
                "data": None
            }
        
        result = read_file(filepath)

        if not result["success"]:
            return result

        content = result["data"]["content"]

        pattern = re.compile(re.escape(keyword),re.IGNORECASE)

        lines = content.splitlines()

        matches = []

        for line_number, line in enumerate(lines, start=1):
            if pattern.search(line):
                matches.append(
                    {
                        "line": line_number,
                        "context": line.strip()
                    }
                )

        return {
            "success": True,
            "message": "Search completed successfully.",
            "data": {
                "filepath": filepath,
                "keyword": keyword,
                "match_count": len(matches),
                "matches": matches
            }
        }
    except Exception as e:
        return {
            "success": False,
            "message": str(e),
            "data": None
        }


